from flask import Flask, request, send_file, render_template
import os, io
from openai import OpenAI
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt
from PyPDF2 import PdfReader
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Load API key
load_dotenv()
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

app = Flask(__name__)

# -------------------------------
# Extract text from file
# -------------------------------
def extract_text(resume_file):
    filename = resume_file.filename.lower()
    text_content = ""

    if filename.endswith(".docx"):
        from docx import Document as DocxReader
        doc = DocxReader(resume_file)
        text_content = "\n".join([p.text for p in doc.paragraphs])

    elif filename.endswith(".pdf"):
        reader = PdfReader(resume_file)
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_content += text + "\n"

    else:
        raise ValueError("Only DOCX and PDF files are supported.")

    return text_content


# -------------------------------
# Build clean DOCX
# -------------------------------
def build_clean_doc(text):
    doc = Document()

    # Base font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    lines = text.split("\n")
    first_line = True

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # -------------------------
        # NAME (first line big + centered)
        # -------------------------
        if first_line:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(20)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            first_line = False
            continue

        # -------------------------
        # CONTACT LINE (centered)
        # -------------------------
        if "@" in line or "|" in line:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            run.font.size = Pt(10)
            continue

        # -------------------------
        # SECTION HEADINGS
        # -------------------------
        if line.upper() in ["PROFILE", "SKILLS", "EXPERIENCE", "EDUCATION"]:
            p = doc.add_paragraph()
            run = p.add_run(line.upper())
            run.bold = True
            run.font.size = Pt(13)

            # Divider line
            doc.add_paragraph("─" * 50)
            continue

        # -------------------------
        # JOB TITLES (bold)
        # -------------------------
        if "–" in line or "|" in line:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(11.5)
            continue

        # -------------------------
        # BULLETS
        # -------------------------
        if line.startswith("- "):
            p = doc.add_paragraph(line[2:], style="List Bullet")
            p.paragraph_format.space_after = Pt(3)
            continue

        # -------------------------
        # NORMAL TEXT
        # -------------------------
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(6)

    return doc


# -------------------------------
# Routes
# -------------------------------
@app.route("/")
def home():
    return render_template("index.html")


@app.route("/optimize_file", methods=["POST"])
def optimize_file():
    resume_file = request.files.get("resume_file")
    job = request.form.get("job")
    ats_keywords = request.form.get("ats_keywords") == "true"

    if not resume_file or not job:
        return "Missing file or job description.", 400

    try:
        text_content = extract_text(resume_file)
    except Exception as e:
        return str(e), 400

    # -------------------------------
    # AI Prompt (PLAIN TEXT ONLY)
    # -------------------------------
    prompt = f"""
You are a professional resume writer.

Rewrite the resume tailored to the job description.

IMPORTANT:
- First line must be FULL NAME
- Second line must be contact info (email | phone | linkedin)
- Keep sections clean and well spaced

RULES:
- Use sections: PROFILE, SKILLS, EXPERIENCE, EDUCATION
- Use bullet points with "-"
- Keep it clean and professional
- Do NOT use markdown (no ##, no **)
- Do NOT return JSON
- Plain text only

FORMAT:

PROFILE
Short paragraph

SKILLS
- Skill 1
- Skill 2

EXPERIENCE
Job Title – Company | Date
- Achievement
- Achievement

EDUCATION
Degree – University | Date
- Detail

Resume:
{text_content}

Job Description:
{job}
"""

    if ats_keywords:
        prompt += "\nInclude relevant ATS keywords naturally."

    try:
        response = client.chat.completions.create(
            model="gpt-4.1-mini",
            messages=[{"role": "user", "content": prompt}],
        )
        output_text = response.choices[0].message.content.strip()

    except Exception as e:
        return f"OpenAI API error: {e}", 500

    # -------------------------------
    # Build DOCX
    # -------------------------------
    doc = build_clean_doc(output_text)

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return send_file(file_stream, as_attachment=True, download_name="Optimized_Resume.docx")


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 10000))  # Use Render's PORT or fallback to 10000
    app.run(host="0.0.0.0", port=port)